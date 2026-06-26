#!/usr/bin/env python3
"""md2docx — Convert markdown to a formatted legal DOCX.

Usage:
    md2docx input.md output.docx
    echo "# Title\n\nBody" | md2docx - output.docx

Markdown conventions for legal documents:
    # TITLE            → Centered, bold, 14pt title
    ## ARTICLE I       → Article heading (bold, 12pt)
    ### Section 1.1    → Section heading (bold, 11pt)
    #### (a)           → Subsection heading
    **bold text**      → Bold (use for defined terms)
    _italic text_      → Italic
    > blockquote       → Indented paragraph (WHEREAS clauses, recitals)
    1. item            → Numbered list
    - item             → Bulleted list
    ---                → Page break
    [SIGNATURE PAGE]   → Page break + centered label
    ___                → Signature line (40 underscores)

The output uses Times New Roman 11pt with legal spacing conventions.
"""

import io
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)

    for level, size in [(1, Pt(14)), (2, Pt(12)), (3, Pt(11)), (4, Pt(11))]:
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = "Times New Roman"
        hstyle.font.size = size
        hstyle.font.bold = True
        hstyle.paragraph_format.space_before = Pt(12)
        hstyle.paragraph_format.space_after = Pt(6)
        if level == 1:
            hstyle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_formatted_runs(paragraph, text: str) -> None:
    """Parse inline markdown (bold, italic) and add runs to the paragraph."""
    parts = re.split(r"(\*\*.*?\*\*|__.*?__|_.*?_|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _heading_level(line: str):
    match = re.match(r"^(#{1,4})\s+", line)
    return len(match.group(1)) if match else None


def _heading_text(line: str) -> str:
    return re.sub(r"^#{1,4}\s+", "", line).strip()


def convert(markdown: str) -> bytes:
    """Convert markdown text to DOCX bytes."""
    doc = Document()
    _setup_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    lines = markdown.split("\n")
    i = 0
    in_blockquote = False
    blockquote_lines: list[str] = []

    def flush_blockquote() -> None:
        nonlocal in_blockquote, blockquote_lines
        if blockquote_lines:
            text = " ".join(blockquote_lines)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            _add_formatted_runs(para, text)
            blockquote_lines = []
        in_blockquote = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            if in_blockquote:
                flush_blockquote()
            i += 1
            continue

        if stripped.startswith(">"):
            content = stripped.lstrip("> ").strip()
            if content:
                in_blockquote = True
                blockquote_lines.append(content)
            i += 1
            continue

        if in_blockquote:
            flush_blockquote()

        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        if stripped.upper().startswith("[SIGNATURE PAGE"):
            doc.add_page_break()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped.strip("[]"))
            run.italic = True
            i += 1
            continue

        if stripped == "___" or stripped.startswith("___"):
            para = doc.add_paragraph()
            para.add_run("_" * 40)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_text = lines[i].strip()
                if not re.match(r"^\|[\s\-:|]+\|$", row_text):
                    table_lines.append(row_text)
                i += 1
            if table_lines:
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.strip("|").split("|")]
                    rows.append(cells)
                if rows:
                    ncols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=ncols)
                    table.style = "Table Grid"
                    for ri, row_cells in enumerate(rows):
                        for ci, cell_text in enumerate(row_cells):
                            if ci < ncols:
                                cell = table.rows[ri].cells[ci]
                                cell.text = ""
                                para = cell.paragraphs[0]
                                if ri == 0:
                                    run = para.add_run(cell_text)
                                    run.bold = True
                                else:
                                    _add_formatted_runs(para, cell_text)
            continue

        level = _heading_level(stripped)
        if level is not None:
            text = _heading_text(stripped)
            para = doc.add_heading(level=level)
            _add_formatted_runs(para, text)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            num = re.match(r"^(\d+)\.", stripped).group(1)
            run = para.add_run(f"{num}. ")
            run.bold = False
            _add_formatted_runs(para, text)
            i += 1
            continue

        if stripped.startswith("- "):
            text = stripped[2:]
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.add_run("• ")
            _add_formatted_runs(para, text)
            i += 1
            continue

        para = doc.add_paragraph()
        _add_formatted_runs(para, stripped)
        i += 1

    if in_blockquote:
        flush_blockquote()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> None:
    if len(sys.argv) < 3:
        print("Usage: md2docx <input.md | -> <output.docx>", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if input_path == "-":
        text = sys.stdin.read()
    else:
        with open(input_path, encoding="utf-8") as f:
            text = f.read()

    docx_bytes = convert(text)

    with open(output_path, "wb") as f:
        f.write(docx_bytes)

    print(f"{output_path} ({len(docx_bytes)} bytes)")


if __name__ == "__main__":
    main()
